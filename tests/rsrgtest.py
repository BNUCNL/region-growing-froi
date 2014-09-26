import time

import numpy as np
import nibabel as nib
from docx import Document
from docx.shared import Inches

from algorithm.regiongrowing import *
from case.matplot_case.histogram_features import show_date_index_formatter


SUBJECT_ID_DIR = "G:/nfs/4Ddata/2006subID"
ACTIVATION_DATA_DIR = "G:/nfs/t2/atlas/group/face-object/activation/2006zstat.nii.gz"
RESULT_NPY_FILE = "peak_points_all_sub.npy"
ROI = ['r_OFA', 'l_OFA', 'r_pFus', 'l_pFus']

RESULT_DATA_DIR = "G:/workingdir/result/"
ASRG_RESULT_DOC_DATA_DIR = "G:/workingdir/result/asrg/doc/"
PC_RESULT_DOCX_FILE = "PC_peak_point_results.docx"
AC_RESULT_DOCX_FILE = "AC_peak_point_results.docx"
PC_RESULT_NPY_FILE = "PC_peak_point_results.npy"
AC_RESULT_NPY_FILE = "AC_peak_point_results.npy"
AC_OPTIMAL_FILE = "AC_optimal_file.nii.gz"
PC_OPTIMAL_FILE = "PC_optimal_file.nii.gz"
TEMP_IMG_DIR = 'temp.png'

if __name__ == "__main__":
    starttime = time.clock()

    mask = nib.load("../data/prior/prob_rFFA.nii.gz")
    mask = mask.get_data()

    seed_coords = np.array(np.nonzero(mask >= 0.6)).T
    neighbor_element = SpatialNeighbor('connected', mask.shape, 26)
    region = Region(seed_coords, neighbor_element)

    similarity_criteria = SimilarityCriteria('euclidean', 0.8)
    stop_criteria = StopCriteria('size')

    image = nib.load(ACTIVATION_DATA_DIR)
    affine = image.get_affine()
    image = image.get_data()

    threshold = np.arange(25, 1525, 25)
    srg = SeededRegionGrowing(similarity_criteria, stop_criteria)

    #save the subject id to lines list
    lines = []
    for line in open(SUBJECT_ID_DIR):
        if line is not '':
            lines.append(line.rstrip('\n'))


    document_PC = Document()
    document_AC = Document()
    document_AC.add_heading(ROI[i] + ' AC Analysis', 0)
    document_PC.add_heading(ROI[i] + ' PC Analysis', 0)

    AC_roi_regions_result = np.zeros((image.shape[3], 1, len(threshold)))
    PC_roi_regions_result = np.zeros((image.shape[3], 1, len(threshold)))

    PC_region_image = np.zeros((image.shape[0], image.shape[1], image.shape[2], image.shape[3]), dtype=int)
    AC_region_image = np.zeros((image.shape[0], image.shape[1], image.shape[2], image.shape[3]), dtype=int)

    # for j in range(0, 1):
    for j in range(0, image.shape[3]):
        seed_sampling_num = 10
        rsrg = RandomSRG(similarity_criteria, stop_criteria, seed_sampling_num)
        rsrg_region = rsrg.compute(region, image, threshold)

        seed_coords = np.array([roi_peak_points[j, i, :]]).astype(np.int)
        neighbor_element = SpatialNeighbor('connected', mask.shape, 26)
        region = Region(seed_coords, neighbor_element)

        # srg_region = srg.compute(region, image[..., j], threshold)
        regions = []
        srg_region = srg.compute(region, image[..., j], threshold)
        regions.append(srg_region)

        #AC
        optimizer_AC = Optimizer('AC')
        optimier_AC_image = optimizer_AC.compute(regions, image[..., j])
        AC_roi_regions_result[j, :] = optimier_AC_image[:]
        title_AC = str(j + 1) + '. ' + lines[j ] + ' --  ' + ROI[i] + ' AC Analysis'
        show_date_index_formatter(threshold, optimier_AC_image[0, :], 'Threshold', 'AC Value', title_AC)

        region_size_AC = threshold[optimier_AC_image.reshape(len(threshold),).argmax()]

        document_AC.add_heading(str(j + 1) + '. ' + lines[j ], 1)
        document_AC.add_paragraph('Peak Point:  ' + str(roi_peak_points[j, i, :]), style='ListBullet')
        document_AC.add_paragraph('Region Optimal Size:  ' + str(region_size_AC), style='ListBullet')
        document_AC.add_paragraph('AC Value MAX:  ' +  str(optimier_AC_image.max()), style='ListBullet')
        document_AC.add_picture(ASRG_RESULT_DOC_DATA_DIR + TEMP_IMG_DIR, width=Inches(4.0))

        #PC
        optimizer_PC = Optimizer('PC')
        optimier_PC_image = optimizer_PC.compute(regions, image[..., j])
        PC_roi_regions_result[j, :] = optimier_PC_image[:]
        title_PC = str(j + 1) + '. ' + lines[j ] + ' --  ' + ROI[i] + ' PC Analysis'
        show_date_index_formatter(threshold, optimier_PC_image[0, :], 'Threshold', 'PC Value', title_PC)

        region_size_PC = threshold[optimier_PC_image.reshape(len(threshold),).argmax()]

        document_PC.add_heading(str(j + 1) + '. ' + lines[j ], 1)
        document_PC.add_paragraph('Peak Point:  ' + str(roi_peak_points[j, i, :]), style='ListBullet')
        document_PC.add_paragraph('Region Optimal Size:  ' + str(region_size_PC), style='ListBullet')
        document_PC.add_paragraph('AC Value MAX:  ' +  str(optimier_PC_image.max()), style='ListBullet')
        document_PC.add_picture(ASRG_RESULT_DOC_DATA_DIR + TEMP_IMG_DIR, width=Inches(4.0))

        #Save data to nii.gz
        AC_labels = srg_region[optimier_AC_image.reshape(len(threshold),).argmax()].get_label()
        AC_region_image[AC_labels[:, 0], AC_labels[:, 1], AC_labels[:, 2], j] = 1
        PC_labels = srg_region[optimier_PC_image.reshape(len(threshold),).argmax()].get_label()
        PC_region_image[PC_labels[:, 0], PC_labels[:, 1], PC_labels[:, 2], j] = 1

        print 'i:', i, '        j: ', j, '       subject_id:', lines[j]

        document_AC.save(ASRG_RESULT_DOC_DATA_DIR + ROI[i] + '_' + AC_RESULT_DOCX_FILE)
        np.save(ASRG_RESULT_DOC_DATA_DIR + ROI[i] + '_' + AC_RESULT_NPY_FILE, AC_roi_regions_result)
        document_AC.save(ASRG_RESULT_DOC_DATA_DIR + ROI[i] + '_' + PC_RESULT_DOCX_FILE)
        np.save(ASRG_RESULT_DOC_DATA_DIR + ROI[i] + '_' + PC_RESULT_NPY_FILE, AC_roi_regions_result)

        nib.save(nib.Nifti1Image(AC_region_image, affine), ASRG_RESULT_DOC_DATA_DIR + ROI[i] + AC_OPTIMAL_FILE)
        nib.save(nib.Nifti1Image(PC_region_image, affine), ASRG_RESULT_DOC_DATA_DIR + ROI[i] + PC_OPTIMAL_FILE)

    print '-----------------------------------------------------------------------------------------'
    read_npy_data = np.load(ASRG_RESULT_DOC_DATA_DIR + 'r_OFA' + '_' + AC_RESULT_NPY_FILE)
    print 'read_npy_data(AC): \n', read_npy_data
    print '-----------------------------------------------------------------------------------------'
    read_npy_data = np.load(ASRG_RESULT_DOC_DATA_DIR + 'r_OFA' + '_' + PC_RESULT_NPY_FILE)
    print 'read_npy_data(PC): \n', read_npy_data

    endtime = time.clock()
    print(endtime - starttime)
    print "Program end..."
































